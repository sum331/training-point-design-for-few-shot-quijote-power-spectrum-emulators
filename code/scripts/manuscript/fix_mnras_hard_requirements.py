"""Apply MNRAS hard-requirement fixes to the z2quijote manuscript draft.

This script intentionally patches both the maintained Markdown source and the
current Word draft.  It addresses issues caught by the internal six-reviewer
pass: figure order, figure alt text, local-path data availability, missing
title-page author metadata, and visible numbering for display equations.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT_DIR = ROOT / "docs" / "manuscript_word"
MD_IN = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260701.md"
MD_OUT = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_hardfix.md"
DOCX_IN = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_review.docx"
DOCX_OUT = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_hardfix.docx"
PDF_OUT = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_hardfix_word_export.pdf"
AUDIT_OUT = MANUSCRIPT_DIR / "z2quijote_mnras_hardfix_audit_20260702.json"


ALT_TEXT: dict[int, str] = {
    1: "Schematic workflow with three paths: Quijote truth for final training and validation, KUN auxiliary spectra for the bias prior, and the design loop producing a 64-point training set.",
    2: "Four-panel schematic showing design-dependent raw bias, standard-geometry acceptance, accepted-only averaging, and interpolation into a continuous bias prior.",
    3: "PPR schematic showing a bias-field potential surface, particle trajectories, repulsion, boundary regularization, and the final cold-start points.",
    4: "Active-selection schematic combining Quijote GP variance and KUN bias proxy into a score optimized inside Delaunay simplex interiors.",
    5: "Bar chart comparing overall p68 relative error for the Sobol64 baseline and the proposed 64-point design.",
    6: "Curve plot of p68 relative error as a function of wavenumber for Sobol64 and the proposed method.",
    7: "Band-wise bar chart comparing p68 relative error across low, mid, focus-high, and tail k-bands.",
    8: "Component-ablation chart showing overall p68 for Sobol32, PPR32, Sobol64, PPR plus Sobol, PPR plus variance-only active selection, PPR plus bias-only active selection, the full method, and Sobol plus variance-bias active selection.",
    9: "Two-panel distribution diagnostic showing empirical error distributions and selected quantiles for the Sobol baseline and the proposed method.",
    10: "Corner-style projection plot comparing Sobol64 baseline points, PPR cold-start points, and active additions in normalized parameter coordinates.",
    11: "Violin and point plot showing the standard-geometry bias-field quantiles occupied by Sobol32, PPR cold-start points, and active additions.",
    12: "Three-panel support diagnostic showing reference-grid support tiers, accepted-count survival, and decreasing bias-estimate uncertainty with accepted count.",
    13: "Two-dimensional slices through the interpolated standard-geometry bias prior across selected normalized parameter pairs.",
}


DATA_AVAILABILITY = """The fixed Quijote validation metadata, final training-design coordinates, ablation summaries, band-wise and quantile source tables, plotting scripts, configuration manifests, and the standard-geometry bias-field support metadata required to reproduce the reported figures will be released in a public repository or journal supplementary package with a stable identifier before external submission. For the present review draft, the same materials are bundled as a local reproducibility package and are described by the accompanying reproducibility manifest. The Quijote simulation suite and the KUN auxiliary emulator are cited in the manuscript; no auxiliary spectra are included in the final Quijote emulator training set."""


@dataclass
class FigureBlock:
    markdown_line: str
    number: int


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_run_font(paragraph: Paragraph, size: float = 9.0, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.italic = italic


def replace_data_availability(md: str) -> str:
    pattern = re.compile(r"## Data availability\n\n.*?\n(?=## References)", re.S)
    replacement = "## Data availability\n\n" + DATA_AVAILABILITY + "\n\n"
    return pattern.sub(replacement, md)


def add_front_matter_metadata(md: str) -> str:
    marker = "\n## Abstract\n"
    if "**Authors:**" in md:
        return md
    metadata = (
        "\n**Authors:** Author list to be finalized before external submission.  \n"
        "**Affiliations:** Institutional affiliations and corresponding-author details to be finalized before external submission.  \n"
        "**Manuscript status:** Internal MNRAS-style review draft, 2 July 2026.\n"
    )
    return md.replace(marker, metadata + marker, 1)


def add_or_replace_alt_text(md: str) -> str:
    lines = md.splitlines()
    out: list[str] = []
    i = 0
    figure_re = re.compile(r"!\[Figure\s+(\d+)\.")
    while i < len(lines):
        line = lines[i]
        out.append(line)
        m = figure_re.search(line)
        if m:
            number = int(m.group(1))
            alt = ALT_TEXT.get(number, "Manuscript figure.")
            if i + 1 < len(lines) and lines[i + 1].startswith("Alt text:"):
                i += 1
            out.append(f"Alt text: {alt}")
        i += 1
    return "\n".join(out) + "\n"


def move_support_figures_in_markdown(md: str) -> str:
    lines = md.splitlines()
    moved: dict[int, str] = {}
    kept: list[str] = []
    figure_line_re = re.compile(r"!\[Figure\s+(12|13)\.")
    skip_next_alt = False
    for line in lines:
        if skip_next_alt and line.startswith("Alt text:"):
            skip_next_alt = False
            continue
        skip_next_alt = False
        m = figure_line_re.search(line)
        if m:
            moved[int(m.group(1))] = line
            skip_next_alt = True
            continue
        kept.append(line)

    md2 = "\n".join(kept) + "\n"
    if "### 5.8 Bias-field support diagnostics" not in md2:
        support_section = "\n".join(
            [
                "### 5.8 Bias-field support diagnostics",
                "",
                "Figure 12 summarizes the usable support of the standard-geometry bias field, including reference-grid coverage, accepted-count survival, and the reduction of estimate scatter with accepted count. Figure 13 provides two-dimensional slices of the interpolated bias prior. These diagnostics are reported after the main validation and mechanism results because they support the construction of the bias prior rather than defining the active-selection outcome.",
                "",
                moved.get(12, ""),
                "",
                moved.get(13, ""),
                "",
            ]
        )
        md2 = md2.replace("### 5.8 Method boundaries and failure modes", support_section + "### 5.9 Method boundaries and failure modes", 1)
    else:
        md2 = md2.replace("### 5.8 Method boundaries and failure modes", "### 5.9 Method boundaries and failure modes")
    return md2


def caption_text_update(md: str) -> str:
    md = md.replace(
        "![Figure 8. Mechanism ablation. The full method gives the lowest overall \\(p68\\). PPR improves the 32-point cold start, bias-only active selection helps, and variance-only or Sobol-seeded active controls do not reproduce the final gain.]",
        "![Figure 8. Mechanism ablation. Lower values are better; the Sobol64 reference is marked for comparison. The full method gives the lowest overall \\(p68\\). PPR improves the 32-point cold start, bias-only active selection helps, and variance-only or Sobol-seeded active controls do not reproduce the final gain.]",
    )
    return md


def number_display_equations(md: str) -> str:
    blocks = list(re.finditer(r"\\\[\n(.*?)\n\\\]", md, re.S))
    if not blocks:
        return md
    out: list[str] = []
    last = 0
    eq_no = 1
    for match in blocks:
        content = match.group(1)
        out.append(md[last : match.start()])
        if "\\tag{" in content:
            out.append(match.group(0))
        else:
            out.append("\\[\n" + content.rstrip() + f"\n\\tag{{{eq_no}}}\n\\]")
            eq_no += 1
        last = match.end()
    out.append(md[last:])
    return "".join(out)


def patch_markdown() -> None:
    md = MD_IN.read_text(encoding="utf-8")
    md = add_front_matter_metadata(md)
    md = move_support_figures_in_markdown(md)
    md = caption_text_update(md)
    md = add_or_replace_alt_text(md)
    md = replace_data_availability(md)
    md = number_display_equations(md)
    MD_OUT.write_text(md, encoding="utf-8")


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="drawing"]'))


def set_drawing_width(paragraph: Paragraph, width_in: float) -> None:
    """Resize the first inline drawing in a paragraph while preserving aspect."""
    extents = paragraph._p.xpath('.//*[local-name()="extent"]')
    if not extents:
        return
    # The paragraph usually contains both wp:extent and a:ext.  The first one
    # is enough to determine the current aspect ratio.
    old_cx = int(extents[0].get("cx"))
    old_cy = int(extents[0].get("cy"))
    if old_cx <= 0 or old_cy <= 0:
        return
    new_cx = int(width_in * 914400)
    new_cy = int(new_cx * old_cy / old_cx)
    for extent in extents:
        extent.set("cx", str(new_cx))
        extent.set("cy", str(new_cy))


def find_caption(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Caption not found: {prefix}")


def previous_paragraph(doc: Document, paragraph: Paragraph) -> Paragraph | None:
    paragraphs = doc.paragraphs
    for idx, candidate in enumerate(paragraphs):
        if candidate._p is paragraph._p:
            return paragraphs[idx - 1] if idx > 0 else None
    return None


def replace_all_text(paragraph: Paragraph, old: str, new: str) -> None:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def patch_docx_text_and_order(doc: Document) -> None:
    # Add title-page metadata after the title if it has not already been added.
    if not any("Affiliations:" in p.text for p in doc.paragraphs[:8]):
        title = doc.paragraphs[0]
        p3 = insert_paragraph_after(title, "Manuscript status: internal MNRAS-style review draft, 2 July 2026", "Normal")
        p2 = insert_paragraph_after(title, "Affiliations: institutional affiliations and corresponding-author details to be finalized before external submission.", "Normal")
        p1 = insert_paragraph_after(title, "Authors: author list to be finalized before external submission.", "Normal")
        for p in [p1, p2, p3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p, size=9)

    # Move Figure 12 and Figure 13 out of Section 3 and place them after Figure 11.
    fig11 = find_caption(doc, "Figure 11.")
    fig12 = find_caption(doc, "Figure 12.")
    fig13 = find_caption(doc, "Figure 13.")
    blocks: list[tuple[Paragraph, Paragraph]] = []
    for cap in [fig12, fig13]:
        image = previous_paragraph(doc, cap)
        if image is None or not paragraph_has_drawing(image):
            raise ValueError(f"Could not find image paragraph before {cap.text[:20]}")
        blocks.append((image, cap))

    anchor = fig11
    intro = insert_paragraph_after(
        anchor,
        "5.8 Bias-field support diagnostics",
        "Heading 2",
    )
    anchor = intro
    desc = insert_paragraph_after(
        anchor,
        "Figures 12 and 13 provide support diagnostics for the standard-geometry bias prior. They are placed after the main validation and mechanism results because they document the quality and spatial structure of the prior rather than define the active-selection outcome.",
        "Normal",
    )
    anchor = desc
    for image, cap in blocks:
        # These diagnostics are now placed inside the two-column result section,
        # so they must be resized to single-column width to avoid clipping.
        set_drawing_width(image, 3.12)
        anchor._p.addnext(image._p)
        anchor = image
        anchor._p.addnext(cap._p)
        anchor = cap

    # Update the following method-boundary heading to preserve the subsection sequence.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5.8 Method boundaries and failure modes":
            paragraph.text = "5.9 Method boundaries and failure modes"

    # Strengthen the ablation caption.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Figure 8. Mechanism ablation."):
            paragraph.text = (
                "Figure 8. Mechanism ablation. Lower values are better; the Sobol64 reference line is marked for comparison. "
                "The full method gives the lowest overall p68. PPR improves the 32-point cold start, bias-only active selection helps, "
                "and variance-only or Sobol-seeded active controls do not reproduce the final gain."
            )
            paragraph.style = "Z2 Caption" if "Z2 Caption" in [s.name for s in doc.styles] else paragraph.style

    # Page-break the band-wise figure away from the p68(k) figure to prevent caption crowding.
    fig7 = find_caption(doc, "Figure 7.")
    fig7_image = previous_paragraph(doc, fig7)
    if fig7_image is not None:
        fig7_image.paragraph_format.page_break_before = True


def replace_data_availability_in_docx(doc: Document) -> None:
    paragraphs = doc.paragraphs
    start = None
    end = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Data availability":
            start = i
        elif start is not None and paragraph.text.strip() == "References":
            end = i
            break
    if start is None or end is None:
        raise ValueError("Data availability section boundaries not found")
    for paragraph in paragraphs[start + 1 : end]:
        paragraph.text = ""
    paragraphs[start + 1].text = DATA_AVAILABILITY
    set_run_font(paragraphs[start + 1], size=10)


def add_alt_text_paragraphs(doc: Document) -> None:
    captions = [p for p in doc.paragraphs if re.match(r"Figure\s+\d+\.", p.text)]
    for cap in captions:
        m = re.match(r"Figure\s+(\d+)\.", cap.text)
        if not m:
            continue
        number = int(m.group(1))
        alt = ALT_TEXT.get(number, "Manuscript figure.")
        nxt = cap._p.getnext()
        if nxt is not None:
            nxt_para = Paragraph(nxt, cap._parent)
            if nxt_para.text.startswith("Alt text:"):
                nxt_para.text = f"Alt text: {alt}"
                set_run_font(nxt_para, size=8, italic=True)
                continue
        p = insert_paragraph_after(cap, f"Alt text: {alt}", "CaptionText" if "CaptionText" in [s.name for s in doc.styles] else "Normal")
        set_run_font(p, size=8, italic=True)
        p.paragraph_format.space_after = Pt(4)


def number_display_equations_in_docx(doc: Document) -> int:
    eq_no = 1
    for paragraph in doc.paragraphs:
        has_omath = bool(paragraph._p.xpath('.//*[local-name()="oMath"]'))
        has_drawing = paragraph_has_drawing(paragraph)
        is_display = has_omath and not has_drawing and not paragraph.text.strip()
        if not is_display:
            continue
        # A right-tab equation number is fragile in this Word two-column draft:
        # Word can push the tabbed number into adjacent text.  Put the visible
        # number on a centered second line instead; the Markdown/LaTeX source
        # still contains explicit \tag{n} markers for the final LaTeX version.
        run = paragraph.add_run(f"\n({eq_no})")
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_no += 1
    return eq_no - 1


def export_pdf(docx: Path, pdf: Path) -> None:
    ps_docx = str(docx.resolve()).replace("'", "''")
    ps_pdf = str(pdf.resolve()).replace("'", "''")
    script = f"""
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open('{ps_docx}', $false, $false)
  $doc.ExportAsFixedFormat('{ps_pdf}', 17)
}} finally {{
  if ($doc -ne $null) {{ $doc.Close($false) | Out-Null }}
  if ($word -ne $null) {{
    $word.Quit() | Out-Null
    [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
  }}
}}
"""
    subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        check=True,
        capture_output=True,
        text=True,
        timeout=240,
    )


def count_pdf_pages(pdf: Path) -> int | None:
    try:
        result = subprocess.run(["pdfinfo", str(pdf)], capture_output=True, text=True, check=False, timeout=30)
        m = re.search(r"Pages:\s+(\d+)", result.stdout)
        return int(m.group(1)) if m else None
    except Exception:
        return None


def text_from_docx(docx: Path) -> str:
    chunks: list[str] = []
    with zipfile.ZipFile(docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    chunks.append(xml)
    return "\n".join(chunks)


def patch_docx() -> dict[str, object]:
    shutil.copy2(DOCX_IN, DOCX_OUT)
    doc = Document(DOCX_OUT)
    patch_docx_text_and_order(doc)
    replace_data_availability_in_docx(doc)
    add_alt_text_paragraphs(doc)
    equation_count = number_display_equations_in_docx(doc)
    doc.save(DOCX_OUT)
    export_pdf(DOCX_OUT, PDF_OUT)
    xml = text_from_docx(DOCX_OUT)
    return {
        "docx": str(DOCX_OUT),
        "pdf": str(PDF_OUT),
        "pdf_pages": count_pdf_pages(PDF_OUT),
        "equations_numbered": equation_count,
        "alt_text_count": xml.count("Alt text:"),
        "local_path_residue_count": len(re.findall(r"[A-Z]:[/\\\\]", xml)),
        "figure_caption_sequence": [
            int(m.group(1))
            for m in re.finditer(r"Figure\s+(\d+)\.", "\n".join(p.text for p in Document(DOCX_OUT).paragraphs))
        ],
        "has_author_placeholder": "Authors:" in "\n".join(p.text for p in Document(DOCX_OUT).paragraphs[:8]),
        "data_availability_has_public_repository": "public repository" in "\n".join(p.text for p in Document(DOCX_OUT).paragraphs),
    }


def main() -> None:
    patch_markdown()
    audit = patch_docx()
    audit.update(
        {
            "markdown": str(MD_OUT),
            "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        }
    )
    AUDIT_OUT.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(audit, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
