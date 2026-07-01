"""Apply the four reviewer-response fixes to the current hardfix DOCX/PDF.

This script works on the already formatted hardfix Word draft to avoid
regenerating the document from an older Markdown source.
"""

from __future__ import annotations

import json
import re
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


PACKAGE_ROOT = Path(__file__).resolve().parents[3]
MANUSCRIPT_DIR = PACKAGE_ROOT / "docs" / "manuscript"
DOCX = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_hardfix.docx"
PDF = MANUSCRIPT_DIR / "z2quijote_full_manuscript_20260702_mnras_hardfix_word_export.pdf"
AUDIT = MANUSCRIPT_DIR / "z2quijote_reviewer_four_fixes_audit_20260702.json"
FIG_KUN = PACKAGE_ROOT / "figures" / "Fig12_kun_proxy_quijote_enrichment.png"


DATA_AVAILABILITY = (
    "The code, processed data products, final training-design coordinates, validation-set metadata, "
    "figure source data, ablation summaries, configuration manifests, and environment reports supporting "
    "this article are available in the public repository "
    "https://github.com/sum331/training-point-design-for-few-shot-quijote-power-spectrum-emulators. "
    "The repository includes REVIEWER_FILES.md, which maps the manuscript results to the corresponding "
    "processed data, plotting scripts, and machine-readable summaries. Internal workstation paths used "
    "during development are retained only in the supplementary reproducibility manifest and are not "
    "required as external access locations. Raw Quijote simulation products and third-party KUN emulator "
    "assets are not redistributed in the repository; they should be obtained from their original providers "
    "subject to their respective data-use conditions. No auxiliary KUN spectra are included in the final "
    "Quijote emulator training set."
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_font(paragraph: Paragraph, size: float = 9.0, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.italic = italic


def find_para(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def replace_para(doc: Document, prefix: str, text: str) -> None:
    paragraph = find_para(doc, prefix)
    paragraph.text = text
    set_font(paragraph, size=9.0)


def add_kun_diagnostic_figure(doc: Document) -> bool:
    if any("KUN bias-prior diagnostic against Quijote cold-start difficulty" in p.text for p in doc.paragraphs):
        return False
    fig11_alt = find_para(doc, "Alt text: Violin and point plot")
    text = (
        "Figure 12 gives a direct diagnostic of the relationship between the KUN standard-geometry bias prior "
        "and Quijote cold-start difficulty. The KUN prior is evaluated at the fixed LHS256 validation coordinates "
        "and ranked by its support-set quantile. The Quijote difficulty measure is the per-validation-point p68 "
        "error of the Sobol32 Quijote emulator, so the comparison uses Quijote truth rather than KUN spectra. "
        "The pointwise Spearman correlation is weak (rho = -0.012), which confirms that the KUN prior should not "
        "be interpreted as a calibrated pointwise Quijote error predictor. Nevertheless, the highest KUN-quantile "
        "region contains a larger fraction of top-quartile Sobol32-error points than the lowest KUN-quantile region "
        "(27.27 percent versus 7.14 percent). In the same high-quantile region, PPR32 produces a positive median "
        "reduction relative to Sobol32. This diagnostic supports the narrower claim used in the paper: the KUN field "
        "provides useful cold-start shape information, not an unbiased or calibrated Quijote error field."
    )
    p = insert_paragraph_after(fig11_alt, text, "Normal")
    set_font(p, size=9.0)

    image_para = insert_paragraph_after(p, "", "Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(FIG_KUN), width=Inches(6.55))

    caption = insert_paragraph_after(
        image_para,
        "Figure 12. KUN bias-prior diagnostic against Quijote cold-start difficulty. "
        "The KUN standard-geometry prior is evaluated on the fixed LHS256 validation coordinates "
        "and compared with Sobol32 held-out Quijote per-sample p68 error and the PPR32 reduction "
        "relative to Sobol32.",
        "Caption",
    )
    set_font(caption, size=8.0)

    alt = insert_paragraph_after(
        caption,
        "Alt text: Three-panel diagnostic showing weak pointwise correlation, mild enrichment of high Sobol32-error "
        "validation points in high KUN-prior quantiles, and positive PPR32 reduction in high-prior regions.",
        "Normal",
    )
    set_font(alt, size=8.0, italic=True)
    return True


def renumber_support_figures(doc: Document) -> None:
    replacements = {
        "Figure 12. Support quality": "Figure 13. Support quality",
        "Figure 13. Two-dimensional slices": "Figure 14. Two-dimensional slices",
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if paragraph.text.startswith(old):
                paragraph.text = paragraph.text.replace(old, new, 1)
                set_font(paragraph, size=8.0)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Figure 12 summarizes the usable support"):
            paragraph.text = paragraph.text.replace("Figure 12 summarizes", "Figure 13 summarizes")
            paragraph.text = paragraph.text.replace("Figure 13 provides", "Figure 14 provides")
            set_font(paragraph, size=9.0)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5.8 Bias-field support diagnostics":
            paragraph.text = "5.8 Bias-field support diagnostics"


def patch_standard_geometry_text(doc: Document) -> None:
    replace_para(
        doc,
        "The discrete support-point estimator is therefore also consistent.",
        "The discrete support-point estimator is therefore also consistent. This is the only unbiasedness statement made in this paper: "
        "the accepted-only Monte Carlo estimator is conditionally unbiased for the standard-geometry conditional mean at the discrete "
        "reference point, given the specified randomized design generator and the event A_SG(x_j,D)=1. When tau_lambda defines a finite "
        "centre target region, the target is the conditional average over that finite standard-geometry event.",
    )
    replace_para(
        doc,
        "This distinction is central to the interpretation of the method.",
        "The continuous field used by PPR is a reliability-weighted interpolation of the discrete support estimates. Interpolation adds "
        "smoothing, boundary, and sparse-support approximation error, so no unbiasedness claim is made for the interpolated field at "
        "arbitrary continuous locations. The target B_SG,F(x) is also a KUN conditional bias prior under standard local geometry, not an "
        "intrinsic Quijote error field and not a pointwise estimate of the final Quijote emulator error. Its role is to provide a "
        "geometry-normalized shape prior for cold-start sample allocation. The scientific claim tested in this paper is that using this "
        "conditional prior improves the Quijote emulator under the fixed validation protocol.",
    )


def patch_ablation_text(doc: Document) -> None:
    replace_para(
        doc,
        "Figure 8 summarizes the completed mechanism ablation",
        "Figure 8 summarizes the mechanism ablation under the same fixed 256-point LHS validation set. The first mechanism is the cold-start "
        "effect. At 32 points, PPR gives p68 = 0.021251, compared with 0.022609 for Sobol32. This 6.01 percent reduction shows that the "
        "standard-geometry bias prior already improves the initial design before any sequential active selection is applied.",
    )
    replace_para(
        doc,
        "The 64-point controls separate this cold-start effect",
        "The second mechanism is the role of the acquisition signal after the PPR initial state is fixed. Filling the remaining budget after "
        "PPR32 with neutral Sobol points gives p68 = 0.019128, close to but slightly worse than Sobol64 (0.019040). PPR32 followed by "
        "variance-only active selection gives p68 = 0.019304, also worse than Sobol64. The variance term alone therefore does not explain "
        "the improvement. PPR32 followed by bias-only active selection reaches p68 = 0.017783, a 6.60 percent reduction relative to Sobol64, "
        "indicating that the KUN bias proxy supplies useful information after the PPR geometry has been established. The full variance-bias "
        "score gives the best result, p68 = 0.014901, reducing Sobol64 by 21.74 percent.",
    )
    replace_para(
        doc,
        "The initial-condition control further clarifies the mechanism.",
        "The third mechanism is the dependence on the initial condition. Starting from Sobol32 and applying the same variance-bias "
        "active-selection rule gives p68 = 0.019858, worse than Sobol64. Thus, the active acquisition rule is not sufficient by itself. "
        "The successful workflow couples a geometry-normalized cold-start prior with sequential acquisition: PPR changes the early simplex "
        "geometry and sample-density layout, while the variance-bias score then uses the current emulator state and the KUN proxy to allocate "
        "the remaining Quijote points. The ablation therefore supports a coupled mechanism: PPR improves cold start, bias-only selection helps "
        "after PPR, variance-only selection is insufficient, and Sobol-seeded variance-bias selection fails to reproduce the final gain.",
    )
    # Remove the old summary paragraph if still present.
    try:
        paragraph = find_para(doc, "These controls support a component-level interpretation.")
        paragraph.text = ""
    except ValueError:
        pass


def patch_discussion_text(doc: Document) -> None:
    replace_para(
        doc,
        "This boundary must be explicit.",
        "This boundary must be explicit. The KUN bias proxy is useful only if its error shape is sufficiently aligned with the regions where "
        "the Quijote emulator is difficult to learn. If the auxiliary spectrum generator responds differently from Quijote in some parameter "
        "regions, B_t(x) may guide selection into regions that do not help Quijote. The present evidence is therefore empirical and component "
        "based. Figure 12 shows that the KUN prior is not a calibrated pointwise predictor of Quijote validation error: the validation-point "
        "Spearman correlation with Sobol32 per-sample p68 is weak. Its useful signal is instead distributional and design-level: high "
        "KUN-prior quantiles are mildly enriched in top-quartile Sobol32-error points, and PPR32 reduces Sobol32 error in the high-prior "
        "region. The mechanism ablation provides the stronger performance evidence. PPR32 improves over Sobol32, bias-only active selection "
        "after PPR32 improves over Sobol64 by 6.60 percent, variance-only selection after PPR32 is worse than Sobol64, and the full "
        "variance-bias score gives the best result. These controls support the use of KUN as a difficult-region shape proxy, while leaving "
        "multi-seed and multi-budget verification as necessary future tests.",
    )
    replace_para(
        doc,
        "Second, the applicability of the KUN bias proxy remains",
        "Second, the applicability of the KUN bias proxy remains an empirical assumption rather than a theorem about Quijote errors. KUN "
        "provides an auxiliary error shape under the same training-point distribution, not a direct observation of Quijote bias. The current "
        "evidence consists of the selected-point bias-quantile diagnostic, the KUN-prior versus Sobol32 cold-start diagnostic, the bias-only "
        "and variance-only active-selection ablations, and the final Quijote validation result. These results show that the proxy is useful "
        "in the present fixed-budget setting, but they do not prove universal agreement between KUN and Quijote residual responses. Future "
        "work should quantify this relationship over multiple seeds, budgets, and validation sets, and should test whether regions highlighted "
        "by KUN also correspond to Quijote validation errors before and after active selection.",
    )


def patch_data_availability(doc: Document) -> None:
    start = None
    end = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Data availability":
            start = i
        elif start is not None and paragraph.text.strip() == "References":
            end = i
            break
    if start is None or end is None:
        raise ValueError("Data availability section not found")
    for paragraph in doc.paragraphs[start + 1 : end]:
        paragraph.text = ""
    doc.paragraphs[start + 1].text = DATA_AVAILABILITY
    set_font(doc.paragraphs[start + 1], size=9.0)


def export_pdf(docx: Path, pdf: Path) -> None:
    ps_docx = str(docx.resolve()).replace("'", "''")
    ps_pdf = str(pdf.resolve()).replace("'", "''")
    script = f"""
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open('{ps_docx}', $false, $false)
  $doc.ExportAsFixedFormat('{ps_pdf}', 17)
}} finally {{
  if ($doc -ne $null) {{ $doc.Close($false) | Out-Null }}
  if ($word -ne $null) {{
    $word.Quit() | Out-Null
    [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
  }}
}}
"""
    subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        check=True,
        capture_output=True,
        text=True,
        timeout=240,
    )


def main() -> None:
    doc = Document(DOCX)
    inserted = add_kun_diagnostic_figure(doc)
    renumber_support_figures(doc)
    patch_standard_geometry_text(doc)
    patch_ablation_text(doc)
    patch_discussion_text(doc)
    patch_data_availability(doc)
    doc.save(DOCX)
    export_pdf(DOCX, PDF)
    text = "\n".join(p.text for p in Document(DOCX).paragraphs)
    audit = {
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "docx": str(DOCX),
        "pdf": str(PDF),
        "inserted_kun_diagnostic": inserted,
        "has_public_repo_data_availability": "https://github.com/sum331/" in text,
        "kun_diagnostic_caption_count": text.count("KUN bias-prior diagnostic against Quijote cold-start difficulty"),
        "figure_numbers": [int(m.group(1)) for m in re.finditer(r"Figure\s+(\d+)\.", text)],
        "local_path_text_hits": len(re.findall(r"[A-Z]:[/\\\\]", text)),
    }
    AUDIT.write_text(json.dumps(audit, indent=2), encoding="utf-8")
    print(json.dumps(audit, indent=2))


if __name__ == "__main__":
    main()
